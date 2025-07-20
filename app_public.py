import streamlit as st
import google.generativeai as genai
from PyPDF2 import PdfReader
from docx import Document

# --- FONCTIONS DE L'APPLICATION ---
def extract_text_from_pdf(pdf_file):
    reader = PdfReader(pdf_file)
    text = ""
    for page in reader.pages:
        text += page.extract_text()
    return text

def extract_text_from_docx(docx_file):
    doc = Document(docx_file)
    text = ""
    for para in doc.paragraphs:
        text += para.text + "\n"
    return text

def analyze_documents(resume_text, job_description):
    # Configure l'API ici pour s'assurer qu'elle est prête pour la requête
    if 'api_key' in st.session_state and st.session_state.api_key:
        genai.configure(api_key=st.session_state.api_key)
    else:
        st.error("La clé API n'est pas configurée.")
        return ""
        
    custom_prompt = f"""
    Analysez le CV suivant en fonction de la description de poste fournie. Vérifiez rigoureusement chaque ligne de la description de poste et évaluez la correspondance exacte avec le CV. 
    Respectez strictement les normes ATS et attribuez des scores uniquement aux éléments corrects. Identifiez :
    1. Le pourcentage de correspondance entre le CV et la description de poste
    2. Une liste précise des mots-clés manquants
    3. Une synthèse finale de l'adéquation globale en 3 points
    4. Des recommandations concrètes pour améliorer le CV avec des exemples

    Présentez les résultats dans l'ordre ci-dessus sans numérotation. 
    Conservez strictement ce format à chaque analyse. 
    Ne modifiez surtout pas la structure du modèle.
    Titre obligatoire : Analyse de CV ATS

    Description de poste : {job_description}
    CV : {resume_text}
    """

    model = genai.GenerativeModel('gemini-1.5-flash')
    response = model.generate_content(custom_prompt)
    return response.text

def rephrase_text(text):
    if 'api_key' in st.session_state and st.session_state.api_key:
        genai.configure(api_key=st.session_state.api_key)
    else:
        st.error("La clé API n'est pas configurée.")
        return ""

    custom_prompt = f"""
    Reformulez le texte suivant selon les normes ATS en ajoutant des mesures quantifiables et des améliorations concrètes :
    Titre obligatoire : Texte reformulé
    Texte original : {text}
    """

    model = genai.GenerativeModel('gemini-1.5-flash')
    response = model.generate_content(custom_prompt)
    return response.text

def display_resume(file):
    file_type = file.name.split('.')[-1].lower()
    if file_type == 'pdf':
        reader = PdfReader(file)
        text = ""
        for page in reader.pages:
            text += page.extract_text()
        st.text_area("Contenu du CV analysé", text, height=400)
    elif file_type == 'docx':
        doc = Document(file)
        text = ""
        for para in doc.paragraphs:
            text += para.text + "\n"
        st.text_area("Contenu du CV analysé", text, height=400)
    else:
        st.error("Format de fichier non supporté. Veuillez uploader un PDF ou DOCX.")

# --- CONFIGURATION DE LA PAGE ---
st.set_page_config(page_title="Système d'Évaluation ATS", layout="wide")

# --- BARRE DE NAVIGATION ---
st.sidebar.title("Navigation")
page = st.sidebar.radio("Aller à", ["Accueil", "Configuration", "Analyseur de CV", "Réécriture ATS", "Modèles ATS"])

# --- ROUTAGE DES PAGES ---
if page == "Accueil":
    st.title("Bienvenue sur l'Analyseur de CV ATS 👋")
    st.markdown("""
    > Bonjour et bienvenue ! Je suis **Kalilou I. Sangare**, professionnel MEAL et passionné de Data Science. 
    > Cette application a été développée pour vous offrir un outil puissant et gratuit d'optimisation de CV. Elle est une adaptation et une amélioration d'un projet initial de **M. V-C-Sai-Santhosh**, pensée pour vous aider à maximiser vos chances auprès des systèmes de suivi des candidats (ATS).
    """)
    
    st.subheader("Comment utiliser cet outil ?")
    st.info("""
    **Étape 1 : Configurez votre clé API**
    - Allez dans l'onglet **Configuration**.
    - Entrez votre clé API personnelle de Google Gemini pour activer l'intelligence artificielle. C'est 100% gratuit et sécurisé.

    **Étape 2 : Analysez votre CV**
    - Rendez-vous sur **Analyseur de CV**.
    - Collez la description du poste que vous visez et téléchargez votre CV.
    - Obtenez instantanément un pourcentage de correspondance, les mots-clés manquants et des conseils d'amélioration.

    **Étape 3 : Optimisez votre contenu**
    - Utilisez la page **Réécriture ATS** pour reformuler des sections de votre CV avec des termes percutants et des données quantifiables.

    **Étape 4 : Inspirez-vous de modèles**
    - Explorez la section **Modèles ATS** pour télécharger des templates de CV déjà optimisés.
    """)
    
    st.markdown("---")
    st.success("Nous espérons que cet outil vous sera précieux dans votre recherche d'emploi. Optimisez votre CV et décrochez le poste de vos rêves !")

elif page == "Configuration":
    st.title("🔑 Configuration de l'API Key")
    st.info("Pour utiliser les fonctionnalités d'analyse, vous avez besoin d'une clé API de Google Gemini. Vous pouvez en obtenir une gratuitement sur [Google AI Studio](https://aistudio.google.com/app/apikey).")
    api_key = st.text_input("Entrez votre clé API Gemini", type="password", help="Votre clé ne sera conservée que pour la durée de votre session actuelle.")
    if st.button("Sauvegarder la clé"):
        if api_key:
            st.session_state.api_key = api_key
            st.success("Clé API sauvegardée avec succès pour cette session !")
        else:
            st.warning("Veuillez entrer une clé API.")

elif page == "Analyseur de CV":
    st.title("📄🔍 Système d'Évaluation ATS")
    st.write("Analysez votre CV par rapport à une description de poste pour optimiser votre candidature.")
    
    if 'api_key' not in st.session_state or not st.session_state.api_key:
        st.warning("Veuillez configurer votre clé API dans la page 'Configuration' pour utiliser cette fonctionnalité.")
    else:
        job_description = st.text_area("Collez ici la description de poste :")
        resume = st.file_uploader("Uploadez votre CV (PDF ou DOCX)", type=["pdf", "docx"])

        if resume:
            st.write("CV uploadé :")
            display_resume(resume)

        if st.button("Analyser la correspondance"):
            if job_description and resume:
                with st.spinner("Analyse en cours..."):
                    resume.seek(0)
                    file_type = resume.name.split('.')[-1].lower()
                    if file_type == 'pdf':
                        resume_text = extract_text_from_pdf(resume)
                    elif file_type == 'docx':
                        resume_text = extract_text_from_docx(resume)
                    
                    analysis = analyze_documents(resume_text, job_description)
                    
                    st.markdown(analysis)
                    
                    if "pourcentage de correspondance" in analysis.lower():
                        match_line = [line for line in analysis.split('\n') if "pourcentage de correspondance" in line.lower()][0]
                        match_percentage = ''.join(filter(str.isdigit, match_line))
                        if match_percentage:
                            st.write(f"Correspondance avec le poste : {match_percentage}%")
                            st.progress(int(match_percentage))

                    st.success("Analyse terminée !")
            else:
                st.error("Veuillez saisir la description de poste et uploader un CV.")

elif page == "Réécriture ATS":
    st.title("🔮 Réécriture ATS")
    st.write("Optimisez vos phrases pour passer les filtres ATS avec des formulations quantifiables.")

    if 'api_key' not in st.session_state or not st.session_state.api_key:
        st.warning("Veuillez configurer votre clé API dans la page 'Configuration' pour utiliser cette fonctionnalité.")
    else:
        text_to_rephrase = st.text_area("Texte à reformuler :")
        
        if st.button("Générer la réécriture"):
            if text_to_rephrase:
                with st.spinner("Réécriture en cours..."):
                    rephrased_text = rephrase_text(text_to_rephrase)
                    st.write(rephrased_text)
                    st.success("Réécriture terminée !")
            else:
                st.error("Veuillez saisir le texte à reformuler.")

elif page == "Modèles ATS":
    st.title("📄📝 Modèles ATS Gratuits")
    st.write("Téléchargez des modèles de CV optimisés pour les systèmes ATS :")

    templates = {
        "Modèle 1": "https://docs.google.com/document/d/1l7Q8m3T-VOaKZ0AtKBMhSe2ui_4vSGCk/edit",
        "Modèle 2": "https://docs.google.com/document/d/18JWABxOw9eQgY1LoGtqPPl3hxHAIuBXK/edit",
        "Modèle 3": "https://docs.google.com/document/d/1D4rXPvf9Z6XXDx3NdfDL8mQxTYmRJxf2/edit",
        "Modèle 4": "https://docs.google.com/document/d/13VMV-3jbdnnIIRXLhM8-xeaRvSW6ItEm/edit",
        "Modèle 5": "https://docs.google.com/document/d/1phlJSdiq2Jf5D4acoPW6GFMMHnElfQuy/edit",
        "Modèle 6": "https://docs.google.com/document/d/1rAuRb4QpVykfSNp1eaJE73qn0MIrlaD8/edit"
    }

    cols = st.columns(3)
    for index, (template_name, template_link) in enumerate(templates.items()):
        col = cols[index % 3]
        col.markdown(f"""
            <div style="text-align:center">
                <iframe src="https://drive.google.com/file/d/{template_link.split('/')[-2]}/preview" width="200" height="250"></iframe>
                <br>
                <a href="{template_link}" target="_blank">{template_name}</a>
            </div>
        """, unsafe_allow_html=True)